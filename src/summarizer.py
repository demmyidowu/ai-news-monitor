import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
import logging

class AINewsSummarizer:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def categorize_articles(self, articles):
        """Enhanced categorization with priority levels"""
        categories = {
            'critical': [],     # Must-read articles
            'research': [],     # Academic papers and studies
            'industry': [],     # Company news and products
            'policy': [],       # Regulation and governance
            'education': []     # Tutorials and guides
        }
        
        for article in articles:
            text = (article['title'] + ' ' + article['summary']).lower()
            source = article.get('source', '').lower()
            priority_level = article.get('priority_level', 'Low')
            
            # Critical articles (high relevance score)
            if priority_level == "Critical":
                categories['critical'].append(article)
            
            # Categorize by content
            elif any(kw in text for kw in ['arxiv', 'research', 'paper', 'study', 'academic']):
                categories['research'].append(article)
            elif any(kw in text for kw in ['regulation', 'policy', 'law', 'ethics', 'governance', 'government']):
                categories['policy'].append(article)
            elif any(kw in text for kw in ['tutorial', 'guide', 'how-to', 'course', 'learning']):
                categories['education'].append(article)
            else:
                # Default to industry news
                categories['industry'].append(article)
        
        return categories
    
    def create_text_summary(self, articles):
        """Create a text-based summary"""
        categories = self.categorize_articles(articles)
        
        summary = f"""# AI News Weekly Summary
Generated on: {datetime.now().strftime('%B %d, %Y')}
Total Articles Analyzed: {len(articles)}

## Executive Summary
This week in AI saw {len(categories['research'])} research developments, {len(categories['industry'])} industry updates, {len(categories['policy'])} policy discussions, and {len(categories['education'])} educational resources.

"""
        
        # Add each category
        category_titles = {
            'critical': '🚨 Critical Developments',
            'research': '🔬 Research & Development',
            'industry': '🏢 Industry News & Products', 
            'policy': '📋 Policy & Ethics',
            'education': '📚 Learning Resources'
        }
        
        for cat_key, title in category_titles.items():
            if categories[cat_key]:
                summary += f"## {title}\n\n"
                # Show more items for critical news
                max_items = 10 if cat_key == 'critical' else 5
                for i, article in enumerate(categories[cat_key][:max_items], 1):
                    priority_indicator = ""
                    if article.get('priority_level') == 'Critical':
                        priority_indicator = " 🔥"
                    elif article.get('priority_level') == 'High':
                        priority_indicator = " ⭐"
                    
                    summary += f"{i}. **{article['title']}{priority_indicator}**\n"
                    summary += f"   Source: {article['source']}\n"
                    summary += f"   Link: {article['link']}\n"
                    if article['summary']:
                        # Truncate summary to first sentence
                        first_sentence = article['summary'].split('.')[0] + '.'
                        summary += f"   Summary: {first_sentence}\n"
                    if article.get('relevance_score'):
                        summary += f"   Relevance Score: {article['relevance_score']}\n"
                    summary += "\n"
        
        return summary
    
    def create_word_document(self, articles, output_path=None, custom_name=None):
        """Create a Word document summary"""
        if output_path is None:
            if custom_name:
                output_path = f"outputs/{custom_name}_{datetime.now().strftime('%Y%m%d')}.docx"
            else:
                output_path = f"outputs/ai_weekly_report_{datetime.now().strftime('%Y%m%d')}.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading('AI News Weekly Summary', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Total Articles Analyzed: {len(articles)}")
        
        categories = self.categorize_articles(articles)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        exec_summary = f"This week in AI saw {len(categories['research'])} research developments, {len(categories['industry'])} industry updates, {len(categories['policy'])} policy discussions, and {len(categories['education'])} educational resources."
        doc.add_paragraph(exec_summary)
        
        # Add each category
        category_info = [
            ('critical', '🚨 Critical Developments'),
            ('research', '🔬 Research & Development'),
            ('industry', '🏢 Industry News'),
            ('policy', '📋 Policy & Ethics'),
            ('education', '📚 Learning Resources')
        ]
        
        for cat_key, title in category_info:
            if categories[cat_key]:
                doc.add_heading(title, level=1)
                
                for i, article in enumerate(categories[cat_key][:5], 1):
                    # Article title
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. {article['title']}").bold = True
                    
                    # Source and link
                    doc.add_paragraph(f"Source: {article['source']}")
                    doc.add_paragraph(f"Link: {article['link']}")
                    
                    # Summary
                    if article['summary']:
                        first_sentence = article['summary'].split('.')[0] + '.'
                        doc.add_paragraph(f"Summary: {first_sentence}")
                    
                    doc.add_paragraph()  # Empty line
        
        doc.save(output_path)
        self.logger.info(f"Word document saved to {output_path}")
        return output_path

# Usage example
if __name__ == "__main__":
    # Load articles from JSON
    with open('data/articles_20250604.json', 'r') as f:
        articles = json.load(f)
    
    summarizer = AINewsSummarizer()
    
    # Create text summary
    text_summary = summarizer.create_text_summary(articles)
    with open('outputs/summary.txt', 'w') as f:
        f.write(text_summary)
    
    # Create Word document
    summarizer.create_word_document(articles)